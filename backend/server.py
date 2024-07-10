from flask import Flask, request, render_template, send_from_directory, jsonify
from werkzeug.utils import secure_filename
import os
from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement

app = Flask(__name__, template_folder='../templates')
UPLOAD_FOLDER = os.path.join(app.root_path, '../static/uploads')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def parse_docx(file_path):
    doc = Document(file_path)
    content = []

    for block in doc.element.body:
        if block.tag.endswith('tbl'):
            table_content = {'type': 'table', 'rows': [], 'borders': []}
            table = block
            for row in table.findall('.//w:tr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                row_content = []
                for cell in row.findall('.//w:tc', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    cell_text = ''.join(cell.find('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}).itertext())
                    row_content.append(cell_text)

                    # Extract cell borders
                    borders = {}
                    tc_pr = cell.find('.//w:tcPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if tc_pr is not None:
                        tc_borders = tc_pr.find('.//w:tcBorders', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if tc_borders is not None:
                            for border_name in ['top', 'left', 'bottom', 'right']:
                                border = tc_borders.find(f'.//w:{border_name}', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                if border is not None:
                                    borders[border_name] = {
                                        'sz': border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz'),
                                        'val': border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'),
                                        'color': border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                                    }
                    table_content['borders'].append(borders)
                table_content['rows'].append(row_content)
            content.append(table_content)
        elif block.tag.endswith('p'):
            para_content = {
                'type': 'paragraph',
                'text': [],
                'bold': [],
                'italic': [],
                'underline': [],
                'font_size': [],
                'alignment': '',
                'is_bullet': False
            }
            para = block
            xml_str = etree.tostring(para).decode()
            xml_tree = etree.fromstring(xml_str.encode('utf-8'))

            alignment = xml_tree.xpath('.//w:pPr/w:jc/@w:val', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            para_content['alignment'] = alignment[0] if alignment else 'left'

            for run in xml_tree.xpath('.//w:r', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                text = ''.join(run.xpath('.//w:t/text()', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
                bold = bool(run.xpath('.//w:b', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
                italic = bool(run.xpath('.//w:i', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
                underline = bool(run.xpath('.//w:u', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
                font_size = run.xpath('.//w:sz/@w:val', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                font_size = int(font_size[0]) / 2 if font_size else None

                para_content['text'].append(text)
                para_content['bold'].append(bold)
                para_content['italic'].append(italic)
                para_content['underline'].append(underline)
                para_content['font_size'].append(font_size)

            is_bullet = xml_tree.xpath('.//w:numPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            para_content['is_bullet'] = bool(is_bullet)

            content.append(para_content)

    return content

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return 'No file part'
    file = request.files['file']
    if file.filename == '':
        return 'No selected file'
    if file:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        content = parse_docx(file_path)
        return jsonify(content)

@app.route('/download', methods=['POST'])
def download_file():
    content = request.json.get('content')
    doc = Document()

    for item in content:
        if item['type'] == 'table':
            table = doc.add_table(rows=len(item['rows']), cols=len(item['rows'][0]))
            for i, row_data in enumerate(item['rows']):
                row = table.rows[i]
                for j, cell_data in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = cell_data

                    # Apply cell borders
                    borders = item['borders'][i * len(row_data) + j]
                    if borders:
                        tc = cell._element
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = OxmlElement('w:tcBorders')
                        for border_name, border_props in borders.items():
                            border_element = OxmlElement(f'w:{border_name}')
                            border_element.set('w:val', border_props['val'])
                            border_element.set('w:sz', border_props['sz'])
                            border_element.set('w:space', '0')
                            border_element.set('w:color', border_props['color'])
                            tcBorders.append(border_element)
                        tcPr.append(tcBorders)
        elif item['type'] == 'paragraph':
            p = doc.add_paragraph()
            if item['alignment'] == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif item['alignment'] == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif item['alignment'] == 'both':
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if item['is_bullet']:
                p.style = 'List Bullet'

            for i, text in enumerate(item['text']):
                run = p.add_run(text)
                if item['bold'][i]:
                    run.bold = True
                if item['italic'][i]:
                    run.italic = True
                if item['underline'][i]:
                    run.underline = True
                if item['font_size'][i]:
                    run.font.size = Pt(item['font_size'][i])

    filename = 'edited.docx'
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    doc.save(file_path)
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
